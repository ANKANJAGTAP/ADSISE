"""
Generate the References document (DOCX) for Expense Tracker PWA.
Lists all libraries, APIs, specifications, and resources used/referred.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(12)

TEAL = RGBColor(0x00, 0x96, 0x88)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = TEAL
    return h

def add_para(text, bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    return p


# ══════════════════════════════════════════════════════════════
# COVER / HEADER
# ══════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()

add_para("Expense Tracker PWA", bold=True, size=28,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
add_para("References & Bibliography", bold=True, size=20,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(20))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─" * 60)
run.font.color.rgb = TEAL

add_para("", size=6)

# Student info
info_lines = [
    "Name: Ankan Jagtap",
    "PRN: 23510058",
    "Topic No: 13",
    "Topic Name: Progressive Web Application (PWA) Technology",
]
for line in info_lines:
    add_para(line, bold=False, size=13, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))

add_para("February 2026", size=13, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# 1. LIBRARIES & PACKAGES USED
# ══════════════════════════════════════════════════════════════
add_heading_styled("1. Libraries & Packages Used", 1)

add_para(
    "The following table lists all npm packages (runtime and development dependencies) "
    "used in this project, along with their versions and official sources."
)

# Table
lib_table = doc.add_table(rows=1, cols=4)
lib_table.style = 'Light Grid Accent 1'
lib_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Package", "Version", "Type", "Official URL"]
for i, h in enumerate(headers):
    cell = lib_table.cell(0, i)
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(11)

libs = [
    ("react", "19.2.0", "Runtime", "https://react.dev"),
    ("react-dom", "19.2.0", "Runtime", "https://react.dev"),
    ("idb", "8.0.3", "Runtime", "https://github.com/jakearchibald/idb"),
    ("vite", "7.3.1", "Dev", "https://vite.dev"),
    ("@vitejs/plugin-react", "5.1.1", "Dev", "https://github.com/vitejs/vite-plugin-react"),
    ("eslint", "9.39.1", "Dev", "https://eslint.org"),
    ("eslint-plugin-react-hooks", "7.0.1", "Dev", "https://www.npmjs.com/package/eslint-plugin-react-hooks"),
    ("eslint-plugin-react-refresh", "0.4.24", "Dev", "https://www.npmjs.com/package/eslint-plugin-react-refresh"),
    ("@eslint/js", "9.39.1", "Dev", "https://www.npmjs.com/package/@eslint/js"),
    ("@types/react", "19.2.7", "Dev", "https://www.npmjs.com/package/@types/react"),
    ("@types/react-dom", "19.2.3", "Dev", "https://www.npmjs.com/package/@types/react-dom"),
    ("globals", "16.5.0", "Dev", "https://www.npmjs.com/package/globals"),
]

for pkg, ver, typ, url in libs:
    row = lib_table.add_row()
    row.cells[0].text = pkg
    row.cells[1].text = ver
    row.cells[2].text = typ
    row.cells[3].text = url
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

add_para("")


# ══════════════════════════════════════════════════════════════
# 2. BROWSER APIs USED
# ══════════════════════════════════════════════════════════════
add_heading_styled("2. Browser APIs & Web Standards Used", 1)

apis = [
    ("IndexedDB API",
     "Client-side NoSQL database built into all modern browsers. Used for persistent "
     "offline storage of expense data.",
     "https://developer.mozilla.org/en-US/docs/Web/API/IndexedDB_API"),
    ("Service Worker API",
     "A background script that intercepts network requests, enabling caching strategies "
     "and offline functionality.",
     "https://developer.mozilla.org/en-US/docs/Web/API/Service_Worker_API"),
    ("Cache API",
     "Used by the Service Worker to store HTTP responses (app shell assets) for offline access.",
     "https://developer.mozilla.org/en-US/docs/Web/API/Cache"),
    ("Web App Manifest (W3C)",
     "A JSON file that describes the PWA to the browser, enabling installability with "
     "custom name, icons, theme color, and display mode.",
     "https://developer.mozilla.org/en-US/docs/Web/Manifest"),
    ("navigator.onLine / online & offline events",
     "Used to detect and reactively display the browser's network connectivity status.",
     "https://developer.mozilla.org/en-US/docs/Web/API/Navigator/onLine"),
    ("Fetch API",
     "Used by the Service Worker to make network requests and clone responses for caching.",
     "https://developer.mozilla.org/en-US/docs/Web/API/Fetch_API"),
]

for name, desc, url in apis:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(12)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    run2 = p2.add_run(desc)
    run2.font.size = Pt(11)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(10)
    run3 = p3.add_run(f"Reference: {url}")
    run3.font.size = Pt(10)
    run3.italic = True
    run3.font.color.rgb = RGBColor(0x00, 0x79, 0x6B)


# ══════════════════════════════════════════════════════════════
# 3. OFFICIAL DOCUMENTATION REFERRED
# ══════════════════════════════════════════════════════════════
add_heading_styled("3. Official Documentation Referred", 1)

docs_referred = [
    ("React Documentation", "https://react.dev/learn",
     "Official React docs — hooks (useState, useEffect), component patterns, StrictMode."),
    ("Vite Documentation", "https://vite.dev/guide/",
     "Vite configuration, plugin setup, development server, and production build."),
    ("idb Library (Jake Archibald)", "https://github.com/jakearchibald/idb",
     "Promise-based IndexedDB wrapper — API reference, usage examples, and migration guides."),
    ("MDN Web Docs — Progressive Web Apps", "https://developer.mozilla.org/en-US/docs/Web/Progressive_web_apps",
     "Comprehensive guide on PWA concepts, Service Workers, manifests, and best practices."),
    ("MDN — Service Worker API", "https://developer.mozilla.org/en-US/docs/Web/API/Service_Worker_API/Using_Service_Workers",
     "Detailed guide on Service Worker lifecycle (install, activate, fetch), caching strategies."),
    ("MDN — IndexedDB", "https://developer.mozilla.org/en-US/docs/Web/API/IndexedDB_API/Using_IndexedDB",
     "Understanding IndexedDB concepts: databases, object stores, transactions, indexes."),
    ("MDN — Web App Manifest", "https://developer.mozilla.org/en-US/docs/Web/Manifest",
     "Manifest properties: name, icons, display, start_url, theme_color, etc."),
    ("Google Developers — PWA", "https://web.dev/explore/progressive-web-apps",
     "Google's PWA guides — installability, offline support, Lighthouse audits."),
    ("Google Developers — Service Workers", "https://developers.google.com/web/fundamentals/primers/service-workers",
     "Introduction to Service Worker lifecycle, caching patterns, and best practices."),
    ("W3C — Web App Manifest Specification", "https://www.w3.org/TR/appmanifest/",
     "The formal W3C specification for the Web Application Manifest."),
]

for title, url, desc in docs_referred:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    run2 = p2.add_run(url)
    run2.font.size = Pt(10)
    run2.italic = True
    run2.font.color.rgb = RGBColor(0x00, 0x79, 0x6B)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(10)
    run3 = p3.add_run(desc)
    run3.font.size = Pt(11)


# ══════════════════════════════════════════════════════════════
# 4. TOOLS & DEVELOPMENT ENVIRONMENT
# ══════════════════════════════════════════════════════════════
add_heading_styled("4. Tools & Development Environment", 1)

tools = [
    ("Node.js", "JavaScript runtime for running Vite, ESLint, and npm scripts.",
     "https://nodejs.org"),
    ("npm", "Node Package Manager — used to install dependencies and run scripts.",
     "https://www.npmjs.com"),
    ("VS Code", "Source code editor used for development.",
     "https://code.visualstudio.com"),
    ("Google Chrome DevTools", "Used for debugging, Lighthouse PWA audits, "
     "Service Worker inspection, IndexedDB viewer, and network throttling.",
     "https://developer.chrome.com/docs/devtools/"),
    ("ESLint", "Static analysis tool for identifying and fixing JavaScript code quality issues.",
     "https://eslint.org"),
    ("Git", "Version control system for tracking code changes.",
     "https://git-scm.com"),
]

for name, desc, url in tools:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{name}: ")
    run.bold = True
    run.font.size = Pt(12)
    run2 = p.add_run(desc)
    run2.font.size = Pt(12)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    run3 = p2.add_run(url)
    run3.font.size = Pt(10)
    run3.italic = True
    run3.font.color.rgb = RGBColor(0x00, 0x79, 0x6B)


# ══════════════════════════════════════════════════════════════
# 5. ACADEMIC & CONCEPTUAL REFERENCES
# ══════════════════════════════════════════════════════════════
add_heading_styled("5. Academic & Conceptual References", 1)

academic = [
    ("Progressive Web Apps: Escaping Tabs Without Losing Our Soul",
     "Alex Russell & Frances Berriman (2015). The original blog post that coined the term 'Progressive Web App' "
     "and outlined the core principles: reliable, fast, engaging.",
     "https://medium.com/@slightlylate/progressive-apps-escaping-tabs-without-losing-our-soul-3b93a8561955"),
    ("Offline Web Applications (W3C)",
     "W3C working group notes on building web applications that function offline using "
     "Service Workers, Cache API, and IndexedDB.",
     "https://www.w3.org/TR/offline-webapps/"),
    ("The Offline Cookbook — Jake Archibald",
     "Comprehensive guide on caching strategies for Service Workers: cache-first, network-first, "
     "stale-while-revalidate, and more.",
     "https://web.dev/articles/offline-cookbook"),
    ("IndexedDB Best Practices",
     "Google Developers guide on structuring IndexedDB databases, handling transactions, "
     "and avoiding common pitfalls.",
     "https://web.dev/articles/indexeddb-best-practices"),
]

for title, desc, url in academic:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    run2 = p2.add_run(desc)
    run2.font.size = Pt(11)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(12)
    run3 = p3.add_run(f"URL: {url}")
    run3.font.size = Pt(10)
    run3.italic = True
    run3.font.color.rgb = RGBColor(0x00, 0x79, 0x6B)


# ══════════════════════════════════════════════════════════════
# 6. SUMMARY TABLE OF ALL REFERENCES
# ══════════════════════════════════════════════════════════════
add_heading_styled("6. Summary Table of All References", 1)

summary_table = doc.add_table(rows=1, cols=3)
summary_table.style = 'Light Grid Accent 1'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(["#", "Reference", "URL"]):
    cell = summary_table.cell(0, i)
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

all_refs = [
    ("React Official Documentation", "https://react.dev"),
    ("Vite Official Documentation", "https://vite.dev"),
    ("idb Library — Jake Archibald", "https://github.com/jakearchibald/idb"),
    ("MDN — Progressive Web Apps", "https://developer.mozilla.org/en-US/docs/Web/Progressive_web_apps"),
    ("MDN — Service Worker API", "https://developer.mozilla.org/en-US/docs/Web/API/Service_Worker_API"),
    ("MDN — IndexedDB API", "https://developer.mozilla.org/en-US/docs/Web/API/IndexedDB_API"),
    ("MDN — Cache API", "https://developer.mozilla.org/en-US/docs/Web/API/Cache"),
    ("MDN — Web App Manifest", "https://developer.mozilla.org/en-US/docs/Web/Manifest"),
    ("MDN — navigator.onLine", "https://developer.mozilla.org/en-US/docs/Web/API/Navigator/onLine"),
    ("MDN — Fetch API", "https://developer.mozilla.org/en-US/docs/Web/API/Fetch_API"),
    ("Google Developers — PWA", "https://web.dev/explore/progressive-web-apps"),
    ("Google — Service Workers Primer", "https://developers.google.com/web/fundamentals/primers/service-workers"),
    ("W3C — Web App Manifest Spec", "https://www.w3.org/TR/appmanifest/"),
    ("The Offline Cookbook — Jake Archibald", "https://web.dev/articles/offline-cookbook"),
    ("IndexedDB Best Practices", "https://web.dev/articles/indexeddb-best-practices"),
    ("ESLint Official Documentation", "https://eslint.org/docs/latest/"),
    ("Node.js Official Documentation", "https://nodejs.org/en/docs/"),
    ("npm Documentation", "https://docs.npmjs.com"),
]

for idx, (name, url) in enumerate(all_refs, 1):
    row = summary_table.add_row()
    row.cells[0].text = str(idx)
    row.cells[1].text = name
    row.cells[2].text = url
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)


# ── Save ──
output_path = "/home/ankan/Progressive Web Application (PWA) Technology/expense-tracker-pwa/Expense_Tracker_PWA_References.docx"
doc.save(output_path)
print(f"✅ References saved to: {output_path}")
