"""
Generate the detailed project report (DOCX) for Expense Tracker PWA.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)

TEAL = RGBColor(0x00, 0x96, 0x88)
DARK = RGBColor(0x21, 0x21, 0x21)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = TEAL
    return h

def add_para(text, bold=False, italic=False, size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # shading
    shading_elm = p.paragraph_format.element.get_or_add_pPr()
    shd = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F5F5F5',
    })
    shading_elm.append(shd)
    return p


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

add_para("Expense Tracker PWA", bold=True, size=28,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
add_para("Progressive Web Application Technology", bold=False, size=18,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(20))

# Horizontal line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─" * 60)
run.font.color.rgb = TEAL
run.font.size = Pt(12)

add_para("", size=8)

# Student info table
table = doc.add_table(rows=4, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ("Name", "Ankan Jagtap"),
    ("PRN", "23510058"),
    ("Topic No.", "13"),
    ("Topic Name", "Progressive Web Application (PWA) Technology"),
]
for i, (label, value) in enumerate(info_data):
    cell_l = table.cell(i, 0)
    cell_r = table.cell(i, 1)
    cell_l.text = label
    cell_r.text = value
    for cell in (cell_l, cell_r):
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(13)
                run.font.name = 'Calibri'
        # Some cell shading for label column
    for paragraph in cell_l.paragraphs:
        for run in paragraph.runs:
            run.bold = True

add_para("", size=8)
add_para("February 2026", bold=False, size=14,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (Manual)
# ══════════════════════════════════════════════════════════════
add_heading_styled("Table of Contents", 1)

toc_items = [
    "1.  Introduction",
    "2.  Problem Statement",
    "3.  Objectives",
    "4.  Technology Stack",
    "5.  System Architecture",
    "6.  Project Structure",
    "7.  Module-wise Description",
    "    7.1  main.jsx — Application Entry Point",
    "    7.2  App.jsx — Root Component",
    "    7.3  ExpenseForm.jsx — Add Expense Form",
    "    7.4  ExpenseList.jsx — Expense List Display",
    "    7.5  db.js — IndexedDB CRUD Module",
    "    7.6  sw.js — Service Worker",
    "    7.7  swRegistration.js — Service Worker Registration",
    "    7.8  manifest.json — PWA Manifest",
    "8.  Key Features",
    "9.  PWA Capabilities Explained",
    "10. UI/UX Design Approach",
    "11. Testing & Deployment",
    "12. Future Enhancements",
    "13. Conclusion",
]
for item in toc_items:
    add_para(item, size=12, space_after=Pt(2))

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════
add_heading_styled("1. Introduction", 1)

add_para(
    "The Expense Tracker PWA is a Progressive Web Application built to help users "
    "track their daily expenses seamlessly, whether they are online or offline. "
    "Unlike traditional web applications that require a constant internet connection, "
    "this application leverages modern web technologies — Service Workers, IndexedDB, "
    "and the Web App Manifest — to deliver a native-app-like experience directly "
    "from the browser."
)
add_para(
    "The application is built using React 19 as the front-end framework, Vite 7 "
    "as the build tool, and the idb library as a Promise-based wrapper around the "
    "browser's IndexedDB API. It is fully client-side — no backend server is required. "
    "All expense data is stored persistently in the user's browser and survives "
    "page refreshes, browser restarts, and offline sessions."
)
add_para(
    "This project is developed as part of Topic No. 13 — Progressive Web Application "
    "(PWA) Technology, to demonstrate the practical implementation of PWA concepts "
    "including offline-first design, installability, and caching strategies."
)


# ══════════════════════════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════
add_heading_styled("2. Problem Statement", 1)

add_para(
    "Managing personal expenses is a common challenge. Many existing solutions either "
    "require an internet connection, involve complex setups, or store sensitive financial "
    "data on remote servers, raising privacy concerns. Users need a simple, lightweight, "
    "and private tool that:"
)
bullets = [
    "Works reliably without an internet connection.",
    "Does not send personal financial data to any server.",
    "Can be installed on any device (mobile, tablet, desktop) without app stores.",
    "Loads instantly, even on slow network connections.",
    "Provides a clean, intuitive interface for adding and reviewing expenses.",
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(12)


# ══════════════════════════════════════════════════════════════
# 3. OBJECTIVES
# ══════════════════════════════════════════════════════════════
add_heading_styled("3. Objectives", 1)

objectives = [
    "Develop a fully functional expense tracking web application using React.",
    "Implement offline-first capability using Service Workers and IndexedDB.",
    "Make the application installable as a PWA on any device.",
    "Use the idb library for clean, Promise-based IndexedDB operations.",
    "Implement a cache-first strategy for static assets and a network-first strategy for dynamic requests.",
    "Provide a responsive, mobile-first UI with a real-time online/offline status indicator.",
    "Ensure all data remains local and private — no backend or cloud dependency.",
]
for obj in objectives:
    p = doc.add_paragraph(obj, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(12)


# ══════════════════════════════════════════════════════════════
# 4. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════
add_heading_styled("4. Technology Stack", 1)

tech_table = doc.add_table(rows=8, cols=3)
tech_table.style = 'Light Grid Accent 1'
tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Technology", "Version", "Purpose"]
for i, h in enumerate(headers):
    cell = tech_table.cell(0, i)
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(11)

tech_data = [
    ("React", "19.2.0", "Component-based UI framework using hooks (useState, useEffect)"),
    ("Vite", "7.3.1", "Fast build tool with Hot Module Replacement (HMR)"),
    ("idb", "8.0.3", "Promise-based wrapper around the IndexedDB API"),
    ("IndexedDB", "Browser API", "Client-side NoSQL database for persistent offline storage"),
    ("Service Worker", "Browser API", "Background script for caching assets and enabling offline access"),
    ("Web App Manifest", "W3C Spec", "JSON file that makes the app installable with custom metadata"),
    ("ESLint", "9.39.1", "Static code analysis tool for consistent code quality"),
]
for row_idx, (tech, ver, purpose) in enumerate(tech_data, start=1):
    tech_table.cell(row_idx, 0).text = tech
    tech_table.cell(row_idx, 1).text = ver
    tech_table.cell(row_idx, 2).text = purpose
    for col in range(3):
        for p in tech_table.cell(row_idx, col).paragraphs:
            for r in p.runs:
                r.font.size = Pt(11)

add_para("")


# ══════════════════════════════════════════════════════════════
# 5. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════
add_heading_styled("5. System Architecture", 1)

add_para(
    "The Expense Tracker PWA follows a layered client-side architecture with no "
    "backend server. All processing and data storage happen within the user's browser."
)

add_para("Architecture Layers:", bold=True, space_after=Pt(4))

layers = [
    ("Presentation Layer (UI)", "React components — App.jsx, ExpenseForm.jsx, ExpenseList.jsx. "
     "These handle the visual interface, form inputs, and expense list rendering."),
    ("State Management Layer", "React hooks (useState, useEffect) manage component state. "
     "The App component tracks the expense list, loading state, and online/offline status."),
    ("Data Persistence Layer", "The db.js module uses the idb library to interact with IndexedDB. "
     "It provides three CRUD operations: addExpense(), getAllExpenses(), deleteExpense()."),
    ("Caching & Offline Layer", "The Service Worker (sw.js) intercepts network requests. "
     "Static assets use a cache-first strategy; dynamic requests use network-first with cache fallback."),
    ("PWA Layer", "The manifest.json file provides app metadata (name, icons, theme color, display mode) "
     "that allows browsers to offer an 'Install' prompt, making the app behave like a native application."),
]
for title, desc in layers:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{title}: ")
    run.bold = True
    run.font.size = Pt(12)
    run2 = p.add_run(desc)
    run2.font.size = Pt(12)

add_para(
    "Data Flow: User interacts with React UI → React calls db.js functions → "
    "db.js reads/writes IndexedDB → React re-renders with updated data. "
    "Simultaneously, the Service Worker caches assets in the background for offline access."
)


# ══════════════════════════════════════════════════════════════
# 6. PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════════
add_heading_styled("6. Project Structure", 1)

tree = """expense-tracker-pwa/
├── public/
│   ├── manifest.json          — PWA configuration (name, icons, theme)
│   ├── sw.js                  — Service Worker (caching, offline)
│   ├── icon-192x192.png       — App icon (192×192)
│   └── icon-512x512.png       — App icon (512×512)
├── src/
│   ├── main.jsx               — Application entry point
│   ├── App.jsx                — Root component (state, handlers)
│   ├── App.css                — Application styles (responsive)
│   ├── index.css              — Global base styles
│   ├── swRegistration.js      — Service Worker registration
│   ├── components/
│   │   ├── ExpenseForm.jsx    — Form to add new expenses
│   │   └── ExpenseList.jsx    — List display with delete & total
│   └── db/
│       └── db.js              — IndexedDB CRUD operations (idb)
├── index.html                 — HTML shell with PWA meta tags
├── vite.config.js             — Vite build configuration
├── package.json               — Dependencies & scripts
└── eslint.config.js           — ESLint configuration"""

add_code_block(tree)


# ══════════════════════════════════════════════════════════════
# 7. MODULE-WISE DESCRIPTION
# ══════════════════════════════════════════════════════════════
add_heading_styled("7. Module-wise Description", 1)

# 7.1 main.jsx
add_heading_styled("7.1  main.jsx — Application Entry Point", 2)
add_para(
    "This is the entry point of the React application. It performs two critical tasks:"
)
add_para("1. Mounts the React component tree into the DOM by rendering <App /> "
         "inside React's <StrictMode> wrapper, targeting the #root div in index.html.")
add_para("2. Calls registerServiceWorker() to register the Service Worker (/sw.js), "
         "enabling PWA capabilities such as caching and offline support.")
add_para("Key imports: StrictMode, createRoot from react-dom/client, App component, "
         "registerServiceWorker function, and index.css for global styles.")

# 7.2 App.jsx
add_heading_styled("7.2  App.jsx — Root Component", 2)
add_para(
    "App.jsx is the root component that orchestrates the entire application. "
    "It manages three pieces of state using the useState hook:"
)
state_items = [
    "expenses — An array of all expense objects loaded from IndexedDB.",
    "isOnline — A boolean tracking the browser's online/offline status (navigator.onLine).",
    "loading — A boolean that is true while expenses are being fetched from IndexedDB on mount.",
]
for s in state_items:
    p = doc.add_paragraph(s, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(12)

add_para(
    "On component mount, a useEffect hook calls getAllExpenses() from db.js to load "
    "persisted data. A second useEffect sets up 'online' and 'offline' event listeners "
    "on the window object to reactively update the isOnline state."
)
add_para(
    "The component provides two handler functions — handleAddExpense() and "
    "handleDeleteExpense() — which call the corresponding db.js functions and then "
    "refresh the expense list by re-reading from IndexedDB."
)
add_para(
    "The render output includes a header with the app title and online/offline badge, "
    "the ExpenseForm component for adding entries, the ExpenseList component for displaying "
    "and deleting entries, and a footer."
)

# 7.3 ExpenseForm.jsx
add_heading_styled("7.3  ExpenseForm.jsx — Add Expense Form", 2)
add_para(
    "ExpenseForm is a controlled React component that renders a form with two input fields:"
)
add_para("• Title (text input) — The name/description of the expense.")
add_para("• Amount (number input) — The expense amount in ₹, with min=0.01 and step=0.01.")
add_para(
    "On form submission, the component trims the title, parses the amount as a float, "
    "and validates both fields. If validation passes, it calls the onAdd callback prop "
    "(provided by App.jsx) with the expense object { title, amount }. After successful "
    "addition, both input fields are reset to empty strings."
)

# 7.4 ExpenseList.jsx
add_heading_styled("7.4  ExpenseList.jsx — Expense List Display", 2)
add_para(
    "ExpenseList receives two props: expenses (array) and onDelete (callback function). "
    "It computes the total of all expenses using Array.reduce() and displays it in a "
    "prominent 'Total Spent' card at the top."
)
add_para(
    "If the expenses array is empty, a friendly message is shown. Otherwise, each expense "
    "is rendered as a list item showing the title, formatted date (using toLocaleDateString "
    "with 'en-IN' locale), amount in ₹, and a delete button (✕) that calls onDelete(id)."
)

# 7.5 db.js
add_heading_styled("7.5  db.js — IndexedDB CRUD Module", 2)
add_para(
    "The db.js module is the data access layer. It uses the idb library (a lightweight "
    "Promise-based wrapper around the raw IndexedDB API) to provide clean async/await "
    "database operations."
)
add_para("Database Configuration:", bold=True)
add_para("• Database Name: ExpenseTrackerDB")
add_para("• Version: 1")
add_para("• Object Store: expenses (keyPath: id, autoIncrement: true)")
add_para("• Index: createdAt (for potential date-based queries)")
add_para("")
add_para("Exported Functions:", bold=True)
funcs = [
    "addExpense(expense) — Creates a new record with title, amount, and an auto-generated "
    "createdAt timestamp (ISO string). Uses a 'readwrite' transaction.",
    "getAllExpenses() — Retrieves all records from the expenses store using a 'readonly' "
    "transaction. Returns an array of expense objects.",
    "deleteExpense(id) — Removes a single record by its auto-incremented id using a "
    "'readwrite' transaction.",
]
for f in funcs:
    p = doc.add_paragraph(f, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(12)

# 7.6 sw.js
add_heading_styled("7.6  sw.js — Service Worker", 2)
add_para(
    "The Service Worker is the backbone of the PWA's offline capability. It operates "
    "as a background script that intercepts all network requests from the application."
)
add_para("Lifecycle Events:", bold=True)
sw_events = [
    "Install Event — Pre-caches the app shell (index.html, manifest.json, icons) into a "
    "versioned static cache. Calls self.skipWaiting() to activate immediately.",
    "Activate Event — Cleans up old caches from previous versions by comparing cache names. "
    "Calls self.clients.claim() to take control of all open tabs.",
    "Fetch Event — Implements a dual caching strategy:\n"
    "  • Cache-first for static assets (app shell) — serves from cache, falls back to network.\n"
    "  • Network-first for dynamic requests — tries network first, caches the response, "
    "falls back to cache if offline.",
]
for s in sw_events:
    p = doc.add_paragraph(s, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(12)

# 7.7 swRegistration.js
add_heading_styled("7.7  swRegistration.js — Service Worker Registration", 2)
add_para(
    "This module exports a single function registerServiceWorker() that handles the "
    "registration of /sw.js. It first checks for browser support ('serviceWorker' in navigator), "
    "then waits for the window 'load' event before registering to avoid competing with "
    "critical resource loading."
)
add_para(
    "It also listens for the 'updatefound' event on the registration object to log when "
    "a new Service Worker version is being installed and activated."
)

# 7.8 manifest.json
add_heading_styled("7.8  manifest.json — PWA Manifest", 2)
add_para(
    "The Web App Manifest is a JSON file that provides metadata about the PWA to the browser. "
    "Key properties include:"
)
manifest_props = [
    "name / short_name — 'Expense Tracker PWA' / 'ExpenseTracker'",
    "description — 'A Progressive Web App to track your daily expenses offline'",
    "start_url — '/' (the app's entry point when launched from home screen)",
    "display — 'standalone' (hides the browser chrome, looks like a native app)",
    "theme_color — '#009688' (teal, matches the app's primary color)",
    "background_color — '#ffffff' (shown during app launch/splash screen)",
    "orientation — 'portrait-primary'",
    "icons — 192×192 and 512×512 PNG icons with 'any' and 'maskable' purposes",
]
for m in manifest_props:
    p = doc.add_paragraph(m, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(12)


# ══════════════════════════════════════════════════════════════
# 8. KEY FEATURES
# ══════════════════════════════════════════════════════════════
add_heading_styled("8. Key Features", 1)

features = [
    ("Offline-First Architecture", "The app works fully without an internet connection. "
     "All data is stored in IndexedDB, and static assets are cached by the Service Worker."),
    ("Installable PWA", "Users can install the app on their device's home screen from the "
     "browser — no app store required. The manifest.json and Service Worker make this possible."),
    ("Add & Delete Expenses", "Users can add new expenses with a title and amount, and delete "
     "any entry with a single click. All operations persist in IndexedDB."),
    ("Real-time Total Calculation", "The total of all expenses is automatically computed and "
     "displayed, updating instantly when expenses are added or removed."),
    ("Online/Offline Status Indicator", "A live badge in the header shows whether the user "
     "is online (🟢) or offline (🔴), using the navigator.onLine API and event listeners."),
    ("Responsive & Mobile-First Design", "The CSS uses flexbox, CSS custom properties, and "
     "media queries to ensure the app looks great on all screen sizes."),
    ("Blazing Fast Load Times", "Vite's optimized build output combined with Service Worker "
     "caching ensures the app loads almost instantly on repeat visits."),
    ("Privacy-First / No Backend", "All data stays in the user's browser. No data is sent "
     "to any server, ensuring complete privacy."),
]
for title, desc in features:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{title}: ")
    run.bold = True
    run.font.size = Pt(12)
    run2 = p.add_run(desc)
    run2.font.size = Pt(12)


# ══════════════════════════════════════════════════════════════
# 9. PWA CAPABILITIES EXPLAINED
# ══════════════════════════════════════════════════════════════
add_heading_styled("9. PWA Capabilities Explained", 1)

add_para(
    "A Progressive Web Application (PWA) is a type of web application that uses modern "
    "web technologies to deliver app-like experiences. The three core pillars of a PWA are:"
)

pillars = [
    ("Reliable", "Loads instantly and works offline or on flaky networks, thanks to Service "
     "Worker caching. In this project, the sw.js pre-caches the app shell during installation "
     "and uses a cache-first strategy for static assets."),
    ("Fast", "Responds quickly to user interactions. Vite's optimized production build, combined "
     "with cached assets, ensures near-instant page loads. React's virtual DOM minimizes "
     "unnecessary re-renders."),
    ("Engaging", "Feels like a native app. The standalone display mode removes the browser chrome. "
     "The app can be installed on the home screen with custom icons and a splash screen. "
     "The theme color integrates with the device's status bar."),
]
for title, desc in pillars:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"{title}: ")
    run.bold = True
    run.font.size = Pt(12)
    run2 = p.add_run(desc)
    run2.font.size = Pt(12)

add_para("Key PWA Technologies Used in This Project:", bold=True, space_after=Pt(4))
pwa_tech = [
    "Service Worker (sw.js) — Caching, offline support, background lifecycle management.",
    "Web App Manifest (manifest.json) — App metadata, installability, icons, display mode.",
    "IndexedDB (via idb) — Client-side persistent storage that works offline.",
    "HTTPS — Required for Service Workers (localhost is exempt during development).",
    "Responsive Viewport — <meta name='viewport'> ensures proper scaling on mobile.",
]
for t in pwa_tech:
    p = doc.add_paragraph(t, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(12)


# ══════════════════════════════════════════════════════════════
# 10. UI/UX DESIGN APPROACH
# ══════════════════════════════════════════════════════════════
add_heading_styled("10. UI/UX Design Approach", 1)

add_para(
    "The application follows a mobile-first, clean, and minimal design philosophy. "
    "Key design decisions include:"
)
design_points = [
    "CSS Custom Properties (Variables) — A centralized theme using --color-primary (#009688), "
    "--color-danger, --color-bg, etc., making it easy to change the entire color scheme.",
    "Flexbox Layout — The app uses flexbox for responsive layouts without complex CSS grid or "
    "third-party layout libraries.",
    "Card-based UI — Expenses are displayed in card-like list items with clear visual hierarchy: "
    "title, date, amount, and delete action.",
    "Online/Offline Badge — A persistent badge in the header provides immediate visual "
    "feedback about network status.",
    "Accessible Forms — Inputs have proper <label> elements, required attributes, "
    "aria-label on delete buttons, and appropriate input types (text, number).",
    "No External CSS Framework — All styles are hand-written in App.css for a lightweight "
    "footprint and full control.",
]
for d in design_points:
    p = doc.add_paragraph(d, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(12)


# ══════════════════════════════════════════════════════════════
# 11. TESTING & DEPLOYMENT
# ══════════════════════════════════════════════════════════════
add_heading_styled("11. Testing & Deployment", 1)

add_para("Development:", bold=True)
add_para("• Run npm run dev to start the Vite development server with HMR.")
add_para("• The app is accessible at http://localhost:5173 by default.")
add_para("• Service Worker registration works on localhost during development.")

add_para("Production Build:", bold=True)
add_para("• Run npm run build to generate optimized static files in the dist/ folder.")
add_para("• Run npm run preview to preview the production build locally.")

add_para("Testing PWA Features:", bold=True)
testing = [
    "Open Chrome DevTools → Application tab → verify manifest, Service Worker, and cache storage.",
    "Use Lighthouse (built into Chrome DevTools) to audit PWA compliance.",
    "Toggle network to offline mode in DevTools to verify the app works without internet.",
    "Test the install prompt by visiting the app in a Chromium-based browser.",
]
for t in testing:
    p = doc.add_paragraph(t, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(12)


# ══════════════════════════════════════════════════════════════
# 12. FUTURE ENHANCEMENTS
# ══════════════════════════════════════════════════════════════
add_heading_styled("12. Future Enhancements", 1)

enhancements = [
    "Edit Expense — Allow users to modify existing expense entries (Update operation).",
    "Category & Tags — Categorize expenses (Food, Transport, Entertainment, etc.) with filter/search.",
    "Charts & Analytics — Visual expense breakdown using charts (e.g., Chart.js or Recharts).",
    "Export to CSV/PDF — Allow users to export their expense data for record keeping.",
    "Cloud Sync — Optional backend sync so expenses can be accessed across devices.",
    "Push Notifications — Remind users to log their expenses daily.",
    "Dark/Light Theme Toggle — Allow users to switch between color themes.",
    "Budget Limits — Set monthly budgets and get warnings when approaching the limit.",
]
for e in enhancements:
    p = doc.add_paragraph(e, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(12)


# ══════════════════════════════════════════════════════════════
# 13. CONCLUSION
# ══════════════════════════════════════════════════════════════
add_heading_styled("13. Conclusion", 1)

add_para(
    "The Expense Tracker PWA successfully demonstrates the power of Progressive Web Application "
    "technology. By combining React for the UI, IndexedDB for persistent client-side storage, "
    "and Service Workers for caching and offline access, the application delivers a fast, "
    "reliable, and engaging user experience — all without requiring a backend server or an "
    "internet connection."
)
add_para(
    "The project showcases key PWA concepts including installability, offline-first design, "
    "cache strategies, and responsive UI — making it a comprehensive, practical example of "
    "modern web development. The fully client-side architecture ensures user privacy, as no "
    "expense data ever leaves the browser."
)
add_para(
    "This project serves as a solid foundation that can be extended with features like "
    "categories, charts, cloud sync, and push notifications to become a full-fledged "
    "personal finance management tool."
)

# ── Save ──
output_path = "/home/ankan/Progressive Web Application (PWA) Technology/expense-tracker-pwa/Expense_Tracker_PWA_Report.docx"
doc.save(output_path)
print(f"✅ Report saved to: {output_path}")
